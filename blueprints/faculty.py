"""Faculty-related routes: login, dashboard, view applications, bills report."""

import logging
from datetime import datetime
from io import BytesIO
from flask import Blueprint, render_template, request, redirect, url_for, session, flash, send_file
from utils.auth_helpers import validate_faculty_credentials
from utils.data_fetcher import get_gf_applications, get_admission_data, get_faculty_bills_data
from utils.faculty_bills_helpers import (
    build_months_structure, 
    filter_and_assign_sl, 
    COMBINED_HEADER
)
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logger = logging.getLogger(__name__)

faculty_bp = Blueprint('faculty', __name__, url_prefix='/faculty')


@faculty_bp.route('/login', methods=['GET', 'POST'])
def login():
    """Faculty login page."""
    error = None
    
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '').strip()
        
        if validate_faculty_credentials(username, password):
            session['logged_in'] = True
            session['role'] = 'Faculty'
            session['username'] = username
            
            logger.info(f"Faculty logged in: {username}")
            flash(f'Welcome, {username}!', 'success')
            return redirect(url_for('faculty.dashboard'))
        else:
            error = 'Invalid username or password. Please try again.'
            logger.warning(f"Failed faculty login attempt: {username}")
    
    return render_template('faculty_login.html', role='Faculty', error=error)


@faculty_bp.route('/dashboard')
def dashboard():
    """Faculty dashboard with links to view applications and data."""
    if session.get('role') != 'Faculty':
        flash('Please log in as faculty', 'warning')
        return redirect(url_for('faculty.login'))
    
    username = session.get('username', 'Faculty')
    return render_template('faculty_dashboard.html', username=username)


@faculty_bp.route('/gfapplications')
def gf_applications():
    """View guest faculty applications (faculty access)."""
    if session.get('role') != 'Faculty':
        flash('Access denied', 'danger')
        return redirect(url_for('faculty.login'))
    
    try:
        data = get_gf_applications()
        if data is None:
            flash('Unable to fetch guest faculty applications', 'warning')
            return render_template('error.html', message='Data temporarily unavailable')
        
        return render_template('gfApp.html', data=data)
    except Exception as e:
        logger.exception("Error displaying GF applications")
        return render_template('error.html', message='Error loading data')


@faculty_bp.route('/admission-applications')
def admission_applications():
    """View student admission applications (faculty access)."""
    if session.get('role') != 'Faculty':
        flash('Access denied', 'danger')
        return redirect(url_for('faculty.login'))
    
    # Reuse admin blueprint logic
    from blueprints.admin import admission_applications as admin_admission_view
    return admin_admission_view()


@faculty_bp.route('/bills-report')
def bills_report():
    """
    Faculty bills monthly/weekly report page.
    Filter by faculty and month, display teaching records with claiming hours calculation.
    """
    if session.get('role') != 'Faculty':
        flash('Access denied', 'danger')
        return redirect(url_for('faculty.login'))
    
    try:
        df = get_faculty_bills_data()
        if df is None or df.empty:
            flash('Unable to fetch faculty bills data', 'warning')
            return render_template('error.html', message='Bills data temporarily unavailable')
        
        months, faculty_list = build_months_structure(df)
        
        # months available sorted newest first
        month_keys = sorted(months.keys(), key=lambda s: datetime.strptime(s, "%B %Y"), reverse=True) if months else []
        
        # selected params
        selected_month = request.args.get("month", month_keys[0] if month_keys else None)
        selected_faculty = request.args.get("faculty", "All")
        
        if selected_month and selected_month not in months:
            selected_month = None
        
        # ensure "All" present in faculty selection
        if "All" not in faculty_list:
            faculty_list = ["All"] + faculty_list
        else:
            faculty_list = sorted(faculty_list)
        
        # default selected faculty to "All" if not provided or invalid
        if selected_faculty not in faculty_list:
            selected_faculty = "All"
        
        rendered_weeks = []
        if selected_month:
            rendered_weeks = filter_and_assign_sl(months, selected_month, selected_faculty)
        
        return render_template("faculty_bills.html",
                               months_list=month_keys,
                               faculty_list=faculty_list,
                               selected_month=selected_month,
                               selected_faculty=selected_faculty,
                               weeks=rendered_weeks,
                               combined_header=COMBINED_HEADER)
    except Exception as e:
        logger.exception("Error loading faculty bills report")
        return render_template('error.html', message='Error loading bills data')


@faculty_bp.route('/bills-report/docx')
def generate_bills_docx():
    """
    Generate DOCX report for faculty bills (?month=...&faculty=...).
    Uses the same filtering & SL assignment logic as bills_report.
    """
    if session.get('role') != 'Faculty':
        flash('Access denied', 'danger')
        return redirect(url_for('faculty.login'))
    
    month = request.args.get("month")
    faculty = request.args.get("faculty", "All")
    
    if not month:
        return "Missing month parameter", 400
    
    try:
        df = get_faculty_bills_data()
        if df is None or df.empty:
            return "Bills data not available", 503
        
        months, faculty_list = build_months_structure(df)
        if month not in months:
            return f"No data for {month}", 404
        
        rendered_weeks = filter_and_assign_sl(months, month, faculty)
        
        if not rendered_weeks:
            return f"No data for {month} and faculty {faculty}", 404
        
        # create DOCX with A4 portrait orientation
        doc = Document()
        
        # Set A4 portrait with narrow margins
        sections = doc.sections
        for section in sections:
            section.page_height = Pt(842)       # A4 height (297mm) 
            section.page_width = Pt(595)        # A4 width (210mm)
            section.left_margin = Pt(36)        # 0.5 inch
            section.right_margin = Pt(36)       # 0.5 inch
            section.top_margin = Pt(36)         # 0.5 inch
            section.bottom_margin = Pt(36)      # 0.5 inch
        
        # Configure default style - Times New Roman 11pt
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)
        
        # Header section matching the image format
        h = doc.add_paragraph()
        run = h.add_run("BANGALORE UNIVERSITY")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        run.underline = True
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        h.space_after = Pt(6)
        
        # Department and Annexure line
        dept_line = doc.add_paragraph()
        dept_run = dept_line.add_run("Department:     BCA")
        dept_run.bold = True
        dept_run.font.size = Pt(11)
        dept_run.font.name = 'Times New Roman'
        dept_run.underline = True
        
        # Add tabs/spaces for annexure
        dept_line.add_run("                    ")
        annex_run = dept_line.add_run("ANNEXURE (time table need to be attached)")
        annex_run.bold = True
        annex_run.font.size = Pt(11)
        annex_run.font.name = 'Times New Roman'
        annex_run.underline = True
        dept_line.space_after = Pt(3)
        
        # Workload line
        workload = doc.add_paragraph()
        workload.add_run("                                                                         ")
        workload_run = workload.add_run("Workload allotted per Week . . . . . .16 hours . . . . . .")
        workload_run.bold = True
        workload_run.font.size = Pt(11)
        workload_run.font.name = 'Times New Roman'
        workload.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        workload.space_after = Pt(8)
        
        # Create tables for each week
        cols = [
            "Sl.\nNo", "Dair\ny No.", "Date", "Particulars / chapter / lectures (as per\nTime Table) I / II / III / IV / V / VI Sem",
            "Actual\nhours", "Claiming hours\n(Lab period\nreduced by 3/4)", "Subject\ncode"
        ]
        
        # Optimized column widths for A4 portrait (595pt width - 72pt margins = ~523pt usable)
        # Squeeze intelligently: minimize fixed columns, maximize content column
        col_widths_pt = [28, 35, 55, 250, 40, 70, 45]  # Total: ~523pt
        
        for w in rendered_weeks:
            week_label = f"Week {w['week_number']}: {w['display_start'].strftime('%d %b %Y')} — {w['display_end'].strftime('%d %b %Y')}"
            week_para = doc.add_paragraph(week_label)
            week_para.space_before = Pt(8)
            week_para.space_after = Pt(4)
            for run in week_para.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
            
            rows = []
            for e in w["entries"]:
                rows.append([
                    str(e.get("SL No", "")),
                    str(e.get("Dairy No.", "")),
                    e.get("Date"),
                    e.get(COMBINED_HEADER, ""),
                    f"{e.get('Actual hours', 0):.1f}",
                    f"{e.get('Claiming hours', 0):.1f}",
                    e.get("Subject code", "")
                ])
            
            if not rows:
                continue
            
            table = doc.add_table(rows=1 + len(rows) + 1, cols=len(cols))
            table.style = 'Table Grid'
            table.autofit = False
            table.allow_autofit = False
            
            # Set column widths in points
            for i, width in enumerate(col_widths_pt):
                for row in table.rows:
                    row.cells[i].width = Pt(width)
            
            # Header row styling - Times New Roman 9pt (smaller for portrait fit)
            hdr_cells = table.rows[0].cells
            for i, c in enumerate(cols):
                p = hdr_cells[i].paragraphs[0]
                p.clear()
                run = p.add_run(c)
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.space_before = Pt(2)
                p.space_after = Pt(2)
            
            # Data rows - Times New Roman 9pt with text wrapping (compact for portrait)
            for r_idx, row_data in enumerate(rows, start=1):
                row_cells = table.rows[r_idx].cells
                for c_idx, val in enumerate(row_data):
                    p = row_cells[c_idx].paragraphs[0]
                    p.clear()
                    run = p.add_run(str(val))
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    p.space_before = Pt(1)
                    p.space_after = Pt(1)
                    # Center align numeric columns
                    if c_idx in [0, 1, 4, 5]:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Enable text wrapping for long content
                    row_cells[c_idx].width = Pt(col_widths_pt[c_idx])
            
            # Totals row - Times New Roman 9pt bold (compact for portrait)
            total_row_cells = table.rows[1 + len(rows)].cells
            for i in range(len(cols)):
                p = total_row_cells[i].paragraphs[0]
                p.clear()
                p.space_before = Pt(2)
                p.space_after = Pt(2)
                
                if i == 0:
                    run = p.add_run("Weekly Total")
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                elif i == 4:
                    run = p.add_run(f"{w['week_total_actual']:.1f}")
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif i == 5:
                    run = p.add_run(f"{w['week_total_claiming']:.1f}")
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Calculate total claiming hours for all weeks
        total_claiming = sum(w['week_total_claiming'] for w in rendered_weeks)
        
        # Add footer section
        doc.add_paragraph("")
        doc.add_paragraph("")
        
        # Total monthly working hours and remuneration
        total_line = doc.add_paragraph()
        total_hours_run = total_line.add_run(f"Total monthly working hours: {total_claiming:.1f}")
        total_hours_run.bold = True
        total_hours_run.font.size = Pt(11)
        total_hours_run.font.name = 'Times New Roman'
        
        total_line.add_run("                    ")
        
        total_renum_run = total_line.add_run(f"Total remuneration claiming - {total_claiming * 1000:.0f}/-")
        total_renum_run.bold = True
        total_renum_run.font.size = Pt(11)
        total_renum_run.font.name = 'Times New Roman'
        total_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        total_line.space_after = Pt(24)
        
        doc.add_paragraph("")
        doc.add_paragraph("")
        
        # Date and Signature line
        date_sig = doc.add_paragraph()
        date_run = date_sig.add_run(f"Date: {datetime.now().strftime('%d / %m / %Y')}")
        date_run.font.size = Pt(11)
        date_run.font.name = 'Times New Roman'
        
        date_sig.add_run("                                                                              ")
        sig_run = date_sig.add_run("Signature of the Guest Faculty")
        sig_run.font.size = Pt(11)
        sig_run.font.name = 'Times New Roman'
        date_sig.space_after = Pt(36)
        
        doc.add_paragraph("")
        doc.add_paragraph("")
        
        # Certification text
        cert = doc.add_paragraph()
        cert.add_run("       ")
        cert_run = cert.add_run(
            "Certified that, the above Guest Faculty has been handled the Classes allotted to him/her as per the Time "
            "Table and as per the Attendance Record maintained in the department. The said dates and hours are is in order."
        )
        cert_run.font.size = Pt(11)
        cert_run.font.name = 'Times New Roman'
        cert.space_after = Pt(48)
        
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("")
        
        # Chairman signature line
        chairman = doc.add_paragraph()
        chairman.add_run("                                                                                                                          ")
        chairman_run = chairman.add_run("Chairman / Chairperson")
        chairman_run.bold = True
        chairman_run.font.size = Pt(11)
        chairman_run.font.name = 'Times New Roman'
        
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        filename = f"report_{month.replace(' ', '_')}_{faculty.replace(' ', '_')}.docx"
        return send_file(bio, as_attachment=True, download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        logger.exception("Error generating DOCX report")
        return "Error generating report", 500


@faculty_bp.route('/logout')
def logout():
    """Log out faculty."""
    session.clear()
    flash('Logged out successfully', 'info')
    return redirect(url_for('home.landing'))
